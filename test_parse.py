# Test DOCX parsing
from docx import Document
import pandas as pd
import os

# file_path = "D:\ProjectAarya\Contract Analyzer\Templates\\agreement-between-manufacturer-and-dealer.docx"
# doc = Document(file_path)
# full_text = "\n".join([para.text for para in doc.paragraphs])

# print("First 500 chars:")
# print(full_text[:1000])
# print(f"\nTotal lines: {len(full_text.splitlines())}")


df1=pd.read_csv("docx_only.csv")

# for filen in df1["source_url"]:
#     doc=Document(filen)
#     full_text="\n".join([para.text for para in doc.paragraphs])
#     with open("files.txt","a") as f:
#         f.write(full_text[:50])

for idx, row in df1.iterrows():
    filename = row["file_name"]  # Use file_name, not source_url
    filepath = row["source_url"]
    
    doc = Document(filepath)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # One file per contract
    output_path = f"data/parsed/{os.path.splitext(filename)[0]}.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(full_text)
    
    print(f"Parsed {filename} ({len(full_text.splitlines())} lines)")

print("Everything printed")