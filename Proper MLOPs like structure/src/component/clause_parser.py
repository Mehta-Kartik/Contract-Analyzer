# import re
# import json
# import os
# from docx import Document
# from src.logger import logging
# from src.config import DataIngestionConfig


# MAIN_CLAUSE_PATTERN = re.compile(r'^\s*(\d+)[\.\)]?\s+(.*)')
# NUMERIC_SUBCLAUSE_PATTERN = re.compile(r'^\s*(\d+\.\d+)[\.\)]?\s+(.*)')
# ALPHA_SUBCLAUSE_PATTERN = re.compile(r'^\s*\(([a-zA-Z0-9]+)\)\s+(.*)')

# STOP_MARKERS = [
#     "IN WITNESS WHEREOF",
#     "DATE:",
#     "PLACE:",
#     "WITNESSESS",
#     "WITNESSES",
#     "PARTY OF FIRST PART",
#     "PARTY OF SECOND PART"
# ]


# def is_stop_line(line: str) -> bool:
#     upper_line = line.upper().strip()
#     return any(marker in upper_line for marker in STOP_MARKERS)


# def parse_clauses(file_path_docx: str, output_path: str) -> str:
#     """
#     Parses a .docx contract file into structured clause JSON.

#     Args:
#         file_path_docx: Absolute or relative path to the .docx contract file.
#         output_path: Path where the structured JSON will be saved.

#     Returns:
#         output_path: The path where JSON was saved.
#     """
#     logging.info("Starting clause parsing.")
#     logging.info(f"Input file: {file_path_docx}")

#     if not os.path.exists(file_path_docx):
#         raise FileNotFoundError(f"Contract file not found: {file_path_docx}")

#     filename = os.path.basename(file_path_docx)

#     # Load DOCX and extract paragraphs directly (no intermediate .txt file)
#     doc = Document(file_path_docx)
#     lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
#     logging.info(f"Extracted {len(lines)} non-empty lines from document.")

#     # Extract agreement heading: first non-empty line before first main clause
#     agreement_heading = None
#     for line in lines:
#         if MAIN_CLAUSE_PATTERN.match(line):
#             break
#         if not agreement_heading:
#             agreement_heading = line
#             break

#     logging.info(f"Agreement heading detected: {agreement_heading}")

#     clauses = []
#     current_clause = None

#     for line in lines:
#         if is_stop_line(line):
#             logging.info(f"Stop marker encountered: '{line}'. Stopping parse.")
#             break

#         numeric_sub_match = NUMERIC_SUBCLAUSE_PATTERN.match(line)
#         alpha_sub_match = ALPHA_SUBCLAUSE_PATTERN.match(line)
#         main_match = MAIN_CLAUSE_PATTERN.match(line)

#         if numeric_sub_match and current_clause is not None:
#             subclause_id = numeric_sub_match.group(1).strip()
#             subclause_text = numeric_sub_match.group(2).strip()
#             current_clause["subclauses"].append({
#                 "subclause_id": subclause_id,
#                 "text": subclause_text
#             })
#             current_clause["clause_text"].append(line)

#         elif alpha_sub_match and current_clause is not None:
#             subclause_id = alpha_sub_match.group(1).strip()
#             subclause_text = alpha_sub_match.group(2).strip()
#             current_clause["subclauses"].append({
#                 "subclause_id": subclause_id,
#                 "text": subclause_text
#             })
#             current_clause["clause_text"].append(line)

#         elif main_match:
#             clause_number = main_match.group(1).strip()
#             clause_title = main_match.group(2).strip()

#             if current_clause is not None:
#                 clauses.append(current_clause)

#             current_clause = {
#                 "agreement_heading": agreement_heading,
#                 "source_file": filename,
#                 "clause_number": clause_number,
#                 "clause_title": clause_title,
#                 "clause_text": [],
#                 "subclauses": []
#             }

#         elif current_clause is not None:
#             current_clause["clause_text"].append(line)

#     if current_clause is not None:
#         clauses.append(current_clause)

#     for clause in clauses:
#         clause["clause_text"] = "\n".join(clause["clause_text"]).strip()

#     final_json = {
#         "agreement_heading": agreement_heading,
#         "source_file": filename,
#         "clauses": clauses
#     }

#     os.makedirs(os.path.dirname(output_path), exist_ok=True)

#     with open(output_path, "w", encoding="utf-8") as f:
#         json.dump(final_json, f, indent=4, ensure_ascii=False)

#     logging.info(f"Saved {len(clauses)} clauses to {output_path}")
#     logging.info(f"Source file recorded: {filename}")
#     return output_path


import re
import json
import os
from typing import Optional, Dict, Any, Tuple, List
from docx import Document
from src.logger import logging
from src.config import DataIngestionConfig


# More forgiving clause pattern: matches 1. 1) 1: 1-
# and also allows 1 with no following text (then text = "")
MAIN_CLAUSE_PATTERN = re.compile(
    r'^\s*(?:clause|section|article)?\s*(\d{1,3})[\.\):\-]?\s*(.*)',
    re.IGNORECASE
)

NUMERIC_SUBCLAUSE_PATTERN = re.compile(
    r'^\s*(\d{1,3}(?:\.\d{1,3})+)[\.\):\-]?\s*(.*)'
)

ALPHA_SUBCLAUSE_PATTERN = re.compile(
    r'^\s*\(([a-zA-Z0-9]{1,5})\)\s*(.*)'
)

ROMAN_SUBCLAUSE_PATTERN = re.compile(
    r'^\s*\(((?:ix|iv|v?i{0,3}|x))\)\s*(.*)',
    re.IGNORECASE
)

STOP_MARKERS = [
    "IN WITNESS WHEREOF",
    "IN WITNESS WHERE OF",
    "DATE:",
    "PLACE:",
    "WITNESSESS",
    "WITNESSES",
    "PARTY OF FIRST PART",
    "PARTY OF SECOND PART",
    "SIGNED AND DELIVERED",
    "SCHEDULE",
    "ANNEXURE"
]


def normalize_text(text: str) -> str:
    return re.sub(r'\s+', ' ', text).strip()


def is_stop_line(line: str) -> bool:
    upper_line = normalize_text(line).upper()
    return any(marker in upper_line for marker in STOP_MARKERS)


def detect_clause_type(line: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    line = normalize_text(line)

    numeric_sub_match = NUMERIC_SUBCLAUSE_PATTERN.match(line)
    if numeric_sub_match:
        return (
            "numeric_sub",
            numeric_sub_match.group(1).strip(),
            normalize_text(numeric_sub_match.group(2) or "")
        )

    # Roman BEFORE alpha to avoid (i), (iv) being misclassified
    roman_sub_match = ROMAN_SUBCLAUSE_PATTERN.match(line)
    if roman_sub_match:
        return (
            "roman_sub",
            roman_sub_match.group(1).strip().lower(),
            normalize_text(roman_sub_match.group(2) or "")
        )

    alpha_sub_match = ALPHA_SUBCLAUSE_PATTERN.match(line)
    if alpha_sub_match:
        return (
            "alpha_sub",
            alpha_sub_match.group(1).strip().lower(),
            normalize_text(alpha_sub_match.group(2) or "")
        )

    main_match = MAIN_CLAUSE_PATTERN.match(line)
    if main_match:
        return (
            "main",
            main_match.group(1).strip(),
            normalize_text(main_match.group(2) or "")
        )

    return None, None, None


def is_likely_new_boundary(line: str) -> bool:
    """
    True only if line starts a real new structural unit.
    """
    line = normalize_text(line)
    clause_type, _, _ = detect_clause_type(line)
    if clause_type is not None:
        return True

    if is_stop_line(line):
        return True

    # If line starts with a digit (1., 2. ...) or clause‑like marker, treat as new boundary
    if re.match(r'^\s*\d', line) or re.match(r'^\s*\(', line):
        return True

    return False


def should_join_with_previous(previous_line: str, current_line: str) -> bool:
    previous_line = normalize_text(previous_line)
    current_line = normalize_text(current_line)

    if not previous_line or not current_line:
        return False

    # NEVER join if current line is a structural boundary
    if is_likely_new_boundary(current_line):
        return False

    if is_stop_line(current_line):
        return False

    # NEVER join if current line starts with a digit (potential clause number)
    if re.match(r'^\d', current_line):
        return False

    # NEVER join if current line looks like a roman/alpha subclause
    if re.match(r'^\([a-zA-Z0-9]{1,5}\)', current_line):
        return False

    if not re.search(r'[.:;]$', previous_line):
        return True

    if current_line and current_line[0].islower():
        return True

    return False


def sentence_split_lines(paragraph_lines: List[str]) -> List[str]:
    """
    Turn each paragraph into one or more sentences, so 1. 2. 3. become separate lines.
    """
    sentences: List[str] = []
    for line in paragraph_lines:
        line = normalize_text(line)
        if not line:
            continue

        # Split at sentence boundaries followed by an uppercase letter or number
        parts = re.split(
            r'(?<=[.!?])\s+(?=[A-Z0-9\(])',
            line
        )
        for part in parts:
            part = part.strip()
            if part:
                sentences.append(part)
    return sentences


def build_logical_lines(paragraph_lines: List[str]) -> List[str]:
    """
    Merge DOCX paragraphs into logical lines, but keep numbered clauses separate.
    """
    # First split long paragraphs into sentences
    sentence_lines = sentence_split_lines(paragraph_lines)

    logical_lines: List[str] = []

    for line in sentence_lines:
        if not line:
            continue

        if not logical_lines:
            logical_lines.append(line)
            continue

        previous_line = logical_lines[-1]

        # Always keep structural lines separate BEFORE should_join_with_previous
        if MAIN_CLAUSE_PATTERN.match(line):
            logical_lines.append(line)
            continue

        if NUMERIC_SUBCLAUSE_PATTERN.match(line):
            logical_lines.append(line)
            continue

        if ROMAN_SUBCLAUSE_PATTERN.match(line) or ALPHA_SUBCLAUSE_PATTERN.match(line):
            logical_lines.append(line)
            continue

        if should_join_with_previous(previous_line, line):
            logical_lines[-1] = f"{previous_line} {line}".strip()
        else:
            logical_lines.append(line)

    return logical_lines


def extract_agreement_heading(raw_lines: List[str]) -> Optional[str]:
    """
    Return the very first non-empty line as the heading.
    """
    for line in raw_lines:
        line = normalize_text(line)
        if line:
            return line
    return None


def finalize_clause(current_clause: Dict[str, Any]) -> Dict[str, Any]:
    if "clause_text" not in current_clause:
        current_clause["clause_text"] = []
    current_clause["clause_text"] = "\n".join(current_clause["clause_text"]).strip()
    return current_clause


def parse_clauses(file_path_docx: str, output_path: str) -> str:
    logging.info("Starting clause parsing.")
    logging.info(f"Input file: {file_path_docx}")

    if not os.path.exists(file_path_docx):
        raise FileNotFoundError(f"Contract file not found: {file_path_docx}")

    filename = os.path.basename(file_path_docx)

    doc = Document(file_path_docx)
    raw_lines = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
    logging.info(f"Extracted {len(raw_lines)} raw non-empty paragraphs from document.")

    # Only the first line counts as the agreement heading
    agreement_heading = extract_agreement_heading(raw_lines)
    logging.info(f"Agreement heading detected: {agreement_heading}")

    # Step 1: split long paragraphs into sentences
    # Step 2: build logical lines that respect clause boundaries
    lines = build_logical_lines(raw_lines)
    logging.info(f"Built {len(lines)} logical lines after sentence splitting and joining.")

    clauses: List[Dict[str, Any]] = []
    current_clause: Optional[Dict[str, Any]] = None
    # print("\n\n\n\n\n\n")
    for line in lines:
        # print("*"*75)
        # print(line)
        if is_stop_line(line):
            logging.info(f"Stop marker encountered: '{line}'. Stopping parse.")
            break

        clause_type, identifier, extracted_text = detect_clause_type(line)

        if clause_type == "main":
            if current_clause is not None:
                clauses.append(finalize_clause(current_clause))

            # Only first sentence is title; rest is clause_text
            full_text = extracted_text or ""
            sentence_split = re.split(
                r'(?<=[.!?])\s+(?=[A-Z])',
                full_text,
                maxsplit=1
            )
            clause_title = sentence_split[0].strip()
            overflow_text = sentence_split[1].strip() if len(sentence_split) > 1 else ""

            current_clause = {
                "agreement_heading": agreement_heading,
                "source_file": filename,
                "clause_number": identifier,
                "clause_title": clause_title,
                "clause_text": [overflow_text] if overflow_text else [],
                "subclauses": []
            }

        elif clause_type in {"numeric_sub", "alpha_sub", "roman_sub"} and current_clause is not None:
            current_clause["subclauses"].append({
                "subclause_id": identifier,
                "text": extracted_text
            })
            current_clause["clause_text"].append(line)

        elif current_clause is not None:
            current_clause["clause_text"].append(line)

    if current_clause is not None:
        clauses.append(finalize_clause(current_clause))

    # Make sure at least an empty clauses list exists
    final_json = {
        "agreement_heading": agreement_heading,
        "source_file": filename,
        "clauses": clauses
    }

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(final_json, f, indent=4, ensure_ascii=False)

    logging.info(f"Saved {len(clauses)} clauses to {output_path}")
    logging.info(f"Source file recorded: {filename}")

    if len(clauses) == 0:
        logging.warning("No clauses were detected. Check numbering format in the document.")

    return output_path
